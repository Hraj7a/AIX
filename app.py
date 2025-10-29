import streamlit as st
import os
from docx import Document
from pypdf import PdfReader
from langdetect import detect
import requests
from azure.ai.translation.text import TextTranslationClient
from azure.core.credentials import AzureKeyCredential
from openai import OpenAI
import io
import time


# ------------------------------------------------------------
# 🔑 CONFIGURATION
# ------------------------------------------------------------

# Azure Translation
AZURE_TRANSLATOR_KEY = os.getenv("AZURE_TRANSLATOR_KEY")
AZURE_TRANSLATOR_ENDPOINT = os.getenv("AZURE_TRANSLATOR_ENDPOINT", "https://salahalali.cognitiveservices.azure.com/")
AZURE_TRANSLATOR_REGION = os.getenv("AZURE_TRANSLATOR_REGION", "qatarcentral")

translator_client = TextTranslationClient(
    endpoint=AZURE_TRANSLATOR_ENDPOINT,
    credential=AzureKeyCredential(AZURE_TRANSLATOR_KEY)
)

# HuggingFace
HF_MODEL_ID = "llmware/industry-bert-contracts-v0.1"
HF_TOKEN = os.getenv("HF_TOKEN")

# OpenAI
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize clients
client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'original_language' not in st.session_state:
    st.session_state.original_language = None
if 'hf_result' not in st.session_state:
    st.session_state.hf_result = None
if 'hf_token' not in st.session_state:
    st.session_state.hf_token = HF_TOKEN or ""

# ------------------------------------------------------------
# 🛠 HELPER FUNCTIONS
# ------------------------------------------------------------
# def get_chatgpt_response(text, country="", model="gpt-4o-mini", hf_analysis=None):
#     """Get analysis from ChatGPT with optional HuggingFace analysis integration."""
#     try:
#         if country:
#                 prompt = f"""Based on the previous legal analysis and the contract text, and considering the laws of {country},
#                 provide a comprehensive analysis incorporating the previous insights and add your analysis on:
#                 1 - Extracted contract information (eg. contracting parties, effective dates, governing laws, financial terms, etc),
#                 2 - Give me details on the missing information from the document if there is any,
#                 3 - Analysis of potential risks, such as non-standard clauses,
#                 4 - Give me legal advice on what to change in the document, opinions, or law comparisons,
#                 5 - Summarized overview of extracted information, missing items, and potential risks.
#                 
#                 Previous Legal Analysis:
#                 {hf_analysis}
#                 
#                 Contract Text:
#                 {text}"""
#         else:
#                 prompt = f"""Based on the previous legal analysis and the contract text, 
#                 provide a comprehensive analysis incorporating the previous insights and add your analysis on:
#                 1 - Extracted contract information (eg. contracting parties, effective dates, governing laws, financial terms, etc),
#                 2 - Give me details on the missing information from the document if there is any,
#                 3 - Analysis of potential risks, such as non-standard clauses,
#                 4 - Give me legal advice on what to change in the document, opinions, or law comparisons,
#                 5 - Summarized overview of extracted information, missing items, and potential risks.
#                 
#                 Previous Legal Analysis:
#                 {hf_analysis}
#                 
#                 Contract Text:
#                 {text}"""

#         messages = [{"role": "user", "content": prompt}]
#         
#         response = client.chat.completions.create(
#             model=model,
#             messages=messages,
#             max_tokens=2048,
#             n=1,
#             temperature=0.7,
#         )
#         return response.choices[0].message.content.strip()
#     except Exception as e:
#         return "Error: Could not get response from ChatGPT"

def translate_text(text, to_language="en"):
    """Translate text to target language using Azure Translator."""
    try:
        response = translator_client.translate(
            body=[{"text": text}],
            to_language=[to_language]
        )
        return response[0].translations[0].text
    except Exception as e:
        st.error(f"Translation error: {e}")
        return text

def extract_text_from_file(file):
    """Extract text from DOCX, PDF, or TXT file."""
    if file.name.endswith(".pdf"):
        pdf = PdfReader(io.BytesIO(file.read()))
        return "\n".join([page.extract_text() or "" for page in pdf.pages]), len(pdf.pages)
    elif file.name.endswith(".docx"):
        doc = Document(io.BytesIO(file.read()))
        return "\n".join([para.text for para in doc.paragraphs]), len(doc.paragraphs) // 20 + 1
    else:
        text = file.read().decode("utf-8", errors="ignore")
        return text, len(text.split('\n'))

def query_huggingface(model_id, token, text, country=""):
    """Send text to the Hugging Face inference API."""
    headers = {"Authorization": f"Bearer {token}"}
    API_URL = f"https://api-inference.huggingface.co/models/{model_id}"
    
    if country:
        prompt = f"""Based on the following text that is taken from a contract document, and based on the laws of {country},
        analyze it for:
        1 - Key contract information and parties
        2 - Missing critical information
        3 - Potential legal risks and non-standard clauses
        4 - Recommendations for improvements
        5 - Overall legal assessment
        Text from legal document: {text}"""
    else:
        prompt = f"""Based on the following text that is taken from a contract document,
        analyze it for:
        0- which parties does this conract fevor more "answer only by it name at first line"?
        1 - Key contract information and parties
        2 - Missing critical information
        3 - Potential legal risks and non-standard clauses
        4 - Recommendations for improvements
        5 - Overall legal assessment
        Text from legal document: {text}"""
    
    payload = {"inputs": prompt, "parameters": {"max_new_tokens": 1024}}
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=120)
        return response
    except Exception as e:
        return None

# ------------------------------------------------------------
# 🎨 UI CONFIGURATION
# ------------------------------------------------------------

def main():
    st.set_page_config(
        page_title="Contract Analysis",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Custom CSS styling
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
        
        * {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
        }
        
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        
        .main {
            background-color: #FAFAFA;
            padding-top: 0 !important;
        }
        
        .stButton>button {
            background: linear-gradient(90deg, #1E8C7E 0%, #2AB598 100%);
            color: white !important;
            border-radius: 12px;
            padding: 0.75rem 2rem;
            font-weight: 600;
            border: none;
            box-shadow: 0 4px 16px rgba(30, 140, 126, 0.3);
            transition: all 0.3s ease;
            font-size: 1rem;
        }
        
        .output-box {
            background: rgba(30, 140, 126, 0.05);
            border: 2px solid #1E8C7E;
            border-radius: 15px;
            padding: 1.5rem;
            margin: 1rem 0;
            min-height: 150px;
        }
        
        .chat-container {
            max-width: 900px;
            margin: 0 auto;
            padding: 2rem 1rem;
        }
        
        .chat-box-user {
            background: #F5F5F5;
            border-radius: 24px;
            padding: 1.5rem 2rem;
            margin-bottom: 1.5rem;
            display: flex;
            align-items: flex-start;
            gap: 1rem;
            max-width: 85%;
            box-shadow: 0 2px 8px rgba(0,0,0,0.04);
            animation: slideInLeft 0.6s ease-out;
        }
        
        .chat-box-ai {
            background: #FFFFFF;
            border-radius: 24px;
            padding: 1.5rem 2rem;
            margin-bottom: 1.5rem;
            margin-left: auto;
            display: flex;
            align-items: flex-start;
            gap: 1rem;
            max-width: 85%;
            box-shadow: 0 4px 16px rgba(0,0,0,0.08);
            animation: slideInRight 0.6s ease-out 0.3s both;
        }
        </style>
    """, unsafe_allow_html=True)

    # Hero section
    st.markdown("""
    <div style="text-align: center; padding: 0.5rem 2rem 2rem 2rem; max-width: 1200px; margin: 0 auto;">
        <h1 style="font-size: 4rem; font-weight: 800; line-height: 1.2; margin-bottom: 1rem; margin-top: 0;">
            <span style="color: #2D3748;">Analyze your Contracts</span><br>
            <span style="color: #A0AEC0;">with</span>
            <span style="color: #1E8C7E;"> Naja7</span>
        </h1>
        <p style="font-size: 1.3rem; color: #718096; max-width: 800px; margin: 0 auto 2rem auto; line-height: 1.6;">
            Upload your PDF and Word documents to extract, analyze, and understand contract content with AI-powered insights.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar configuration
    st.sidebar.markdown("## Settings")
    hf_model_id_input = st.sidebar.text_input("HF Model ID", value=HF_MODEL_ID)
    st.sidebar.caption("Change this if you get a 404: verify the exact model repo name.")

    # HF Token control
    hf_token_input = st.sidebar.text_input("HF Token", value=st.session_state.hf_token, type="password")
    if hf_token_input != st.session_state.hf_token:
        st.session_state.hf_token = hf_token_input
    if st.session_state.hf_token:
        st.sidebar.success("HF token detected and will be used for requests.")
    else:
        st.sidebar.warning("No HF token set; HF analysis will be skipped.")

    # Main content area
    col1, col2 = st.columns([2, 1])

    with col1:
        # File uploader section
        with st.form("analysis_form", clear_on_submit=False):
            uploaded_file = st.file_uploader(
            "Upload Contract Document",
            type=["txt", "pdf", "docx"],
            help="Supported formats: PDF (.pdf), Word (.docx), and text (.txt) files"
            )
            country = st.text_input("Specify country/region (Optional)", "")
            submitted = st.form_submit_button("Analyze")
        if uploaded_file:
            # Step 1: Process uploaded file and handle translation
            with st.spinner("Processing document..."):
                text, num_pages = extract_text_from_file(uploaded_file)
                
                if not text.strip():
                    st.error(f"The file {uploaded_file.name} appears to be empty or unreadable.")
                    return

                # Detect language
                try:
                    lang = detect(text[:500])
                    st.session_state.original_language = lang
                    st.info(f"Detected language: {lang.upper()}")
                except:
                    lang = "unknown"
                    st.session_state.original_language = "en"
                    st.warning("Could not detect language, proceeding with analysis...")

                # Translate if Arabic
                if lang == "ar":
                    with st.spinner("Translating Arabic text to English..."):
                        try:
                            text = translate_text(text, to_language="en")
                        except Exception as e:
                            st.error(f"Translation error: {e}")
                            return

            # Display document content
            with st.expander("📄 Document Content", expanded=False):
                st.text_area("Content", text, height=200)

            # Step 2: Process with HuggingFace model
            with st.spinner("Performing initial legal analysis..."):
                st.session_state.hf_result = None
                if not st.session_state.hf_token:
                    st.warning("Hugging Face token is not set. Please configure HF_TOKEN to use the HF model.")
                else:
                    try:
                        hf_response = query_huggingface(hf_model_id_input, st.session_state.hf_token, text, country)
                        st.session_state.hf_result = None

                        # Handle no response (network/timeout)
                        if hf_response is None:
                            st.warning("Could not reach Hugging Face Inference API (network/timeout).")
                        else:
                            # Handle model loading (503) with limited retries
                            if hf_response.status_code == 503:
                                try:
                                    payload = hf_response.json()
                                except Exception:
                                    payload = {}
                                estimated = payload.get("estimated_time", 10)
                                st.info(f"HF model is loading. Retrying in {int(estimated)}s...")
                                time.sleep(min(int(estimated), 15))
                                # Retry a couple of times
                                for _ in range(2):
                                    retry_resp = query_huggingface(hf_model_id_input, st.session_state.hf_token, text, country)
                                    if retry_resp is not None and retry_resp.status_code == 200:
                                        hf_response = retry_resp
                                        break
                                    time.sleep(2)

                            # If OK, parse output
                            if hf_response is not None and hf_response.status_code == 200:
                                try:
                                    data = hf_response.json()
                                    if isinstance(data, dict) and "generated_text" in data:
                                        st.session_state.hf_result = data["generated_text"]
                                    elif isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict) and "generated_text" in data[0]:
                                        st.session_state.hf_result = data[0]["generated_text"]
                                    else:
                                        st.warning("HF response received but in an unexpected format. Showing raw output below.")
                                        st.code(str(data))
                                except Exception as parse_err:
                                    st.error(f"Failed to parse HF response: {parse_err}")
                            elif hf_response is not None:
                                # Show precise error for common statuses
                                try:
                                    err_text = hf_response.text
                                except Exception:
                                    err_text = ""
                                if hf_response.status_code in (401, 403):
                                    st.error("Hugging Face authorization failed (check HF_TOKEN permissions or model visibility).")
                                elif hf_response.status_code == 404:
                                    st.error("HF model not found (404). Please verify the model ID in the sidebar.")
                                else:
                                    st.warning(f"HF model request failed: {hf_response.status_code}. Details: {err_text[:500]}")
                    except Exception as e:
                        st.warning(f"Initial analysis error: {str(e)}")

            # Step 3: Process with GPT using HF insights
            with st.spinner("Performing comprehensive analysis..."):
                if st.session_state.hf_result:
                    # ✅ Use only Hugging Face analysis (no GPT)
                    st.session_state.analysis_result = st.session_state.hf_result


            st.markdown("###Analysis Results")
            st.write(st.session_state.analysis_result)

            # Translate back to Arabic if needed
            if lang == "ar":
                with st.spinner("Translating analysis to Arabic..."):
                    try:
                        arabic_response = translate_text(st.session_state.analysis_result, to_language="ar")
                        st.markdown("### 🔄 Arabic Translation")
                        st.markdown('<div class="output-box">', unsafe_allow_html=True)
                        st.write(arabic_response)
                        st.markdown('</div>', unsafe_allow_html=True)
                    except Exception as e:
                        st.error(f"Translation error: {e}")

    # Chat interface in the second column
    with col2:
        st.markdown("### 💬 Chat with AI")
        
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])

        # Chat input
# Chat input + manual submit option
    chat_col = st.container()
    with chat_col:
        user_question = st.text_input("Ask about the document...", key="chat_question")
        send_button = st.button("Send")

    if send_button and user_question:
        st.session_state.messages.append({"role": "user", "content": user_question})

    # Create context from analysis
        context = f"""Based on this analysis of the legal document:
        {st.session_state.analysis_result}
    
        Answer this question: {user_question}"""

    # Get response from GPT-4o-mini
        response = get_chatgpt_response(context, model="gpt-4o-mini")

    # Translate back if needed
        if st.session_state.original_language == "ar":
            try:
                response = translate_text(response, to_language="ar")
            except Exception as e:
                st.error(f"Translation error: {e}")

    # Store assistant message
        st.session_state.messages.append({"role": "assistant", "content": response})

    # (No duplicate rendering here.) Messages are displayed above in the chat column.


if __name__ == "__main__":
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "analysis_result" not in st.session_state:
        st.session_state.analysis_result = ""
    if "hf_result" not in st.session_state:
        st.session_state.hf_result = None
    main()
