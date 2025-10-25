import streamlit as st
import pandas as pd
import numpy as np
import PyPDF2
from io import BytesIO
import re
from docx import Document

# Set page configuration
st.set_page_config(
    page_title="AI Document Analyzer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS inspired by V7 Labs design
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    /* Global styles */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Main container - no top padding */
    .main {
        background-color: #FAFAFA;
        padding-top: 0 !important;
    }
    
    /* V7-style hero section */
    .v7-hero {
        text-align: center;
        padding: 4rem 2rem;
        max-width: 1200px;
        margin: 0 auto;
    }
    
    .v7-logo {
        font-size: 2rem;
        font-weight: 700;
        color: #1E8C7E;
        margin-bottom: 3rem;
        letter-spacing: -0.5px;
    }
    
    .v7-headline {
        font-size: 4.5rem;
        font-weight: 800;
        line-height: 1.1;
        color: #000000;
        margin-bottom: 1.5rem;
        letter-spacing: -2px;
    }
    
    .v7-subheadline {
        font-size: 4.5rem;
        font-weight: 800;
        line-height: 1.1;
        color: #000000;
        margin-bottom: 4rem;
        letter-spacing: -2px;
    }
    
    .highlight-green {
        color: #1E8C7E;
        background: linear-gradient(135deg, #1E8C7E 0%, #2AB598 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    /* Chat boxes container */
    .chat-container {
        max-width: 900px;
        margin: 0 auto;
        padding: 2rem 1rem;
    }
    
    /* User message box (left) */
    .chat-box-user {
        background: #F5F5F5;
        border-radius: 24px;
        padding: 1.5rem 2rem;
        margin-bottom: 1.5rem;
        display: flex;
        align-items: flex-start;
        gap: 1rem;
        max-width: 85%;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        animation: slideInLeft 0.6s ease-out;
    }
    
    /* AI response box (right) */
    .chat-box-ai {
        background: #FFFFFF;
        border-radius: 24px;
        padding: 1.5rem 2rem;
        margin-bottom: 1.5rem;
        margin-left: auto;
        display: flex;
        align-items: flex-start;
        gap: 1rem;
        max-width: 85%;
        box-shadow: 0 4px 16px rgba(0,0,0,0.08);
        animation: slideInRight 0.6s ease-out 0.3s both;
    }
    
    .avatar {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        flex-shrink: 0;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 600;
        font-size: 0.9rem;
    }
    
    .avatar-user {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    .avatar-ai {
        background: linear-gradient(135deg, #1E8C7E 0%, #2AB598 100%);
        color: white;
    }
    
    .chat-content {
        flex: 1;
        padding-top: 0.2rem;
    }
    
    .chat-text {
        font-size: 1rem;
        line-height: 1.6;
        color: #333333;
        margin: 0;
    }
    
    /* Animations */
    @keyframes slideInLeft {
        from {
            opacity: 0;
            transform: translateX(-30px);
        }
        to {
            opacity: 1;
            transform: translateX(0);
        }
    }
    
    @keyframes slideInRight {
        from {
            opacity: 0;
            transform: translateX(30px);
        }
        to {
            opacity: 1;
            transform: translateX(0);
        }
    }
    
    /* Sidebar styling: dark theme to match main content (remove white area) */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0F1720 0%, #0B0F13 100%);
        color: #cbd5e1; /* light text for contrast */
        border-right: none;
        padding: 1rem 1.25rem;
        min-width: 240px;
    }

    /* Make all text in sidebar more legible on dark background */
    [data-testid="stSidebar"] * {
        color: #cbd5e1 !important;
    }

    /* Adjust expander and buttons inside sidebar for dark theme */
    [data-testid="stSidebar"] .stExpander > div {
        background: rgba(255,255,255,0.02) !important;
        border: 1px solid rgba(255,255,255,0.03) !important;
        color: #cbd5e1 !important;
    }

    [data-testid="stSidebar"] .stButton>button {
        background: rgba(255,255,255,0.03) !important;
        color: #cbd5e1 !important;
        border: 1px solid rgba(255,255,255,0.04) !important;
    }
    
    /* Navigation styling */
    .stRadio > label {
        font-weight: 600;
        font-size: 0.95rem;
        color: #333333;
    }
    
    .stRadio > div {
        gap: 0.5rem;
    }
    
    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding: 10px 24px;
        background-color: rgba(240, 242, 246, 0.1);
        border-radius: 10px;
        font-weight: 600;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #1E8C7E 0%, #2AB598 100%);
        color: white !important;
    }
    
    /* Feature cards */
    .feature-card {
        background: rgba(255, 255, 255, 0.05);
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.08);
        margin: 1rem 0;
        border-left: 4px solid #1E8C7E;
    }
    .feature-card h3, .feature-card h4 {
        color: #1E8C7E !important;
    }
    
    /* Buttons */
    .stButton>button {
        background: linear-gradient(90deg, #1E8C7E 0%, #2AB598 100%);
        color: white !important;
        border-radius: 12px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        border: none;
        box-shadow: 0 4px 16px rgba(30, 140, 126, 0.3);
        transition: all 0.3s ease;
        font-size: 1rem;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 24px rgba(30, 140, 126, 0.4);
    }
    
    /* Output boxes styling */
    .output-box {
        background: rgba(30, 140, 126, 0.05);
        border: 2px solid #1E8C7E;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        min-height: 150px;
    }
    
    @media (prefers-color-scheme: dark) {
        .output-box {
            background: rgba(30, 140, 126, 0.1);
            border-color: #2AB598;
        }
    }
    
    /* Q&A Box styling */
    .qa-box {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
        border: 2px solid #667eea;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    @media (prefers-color-scheme: dark) {
        .qa-box {
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.15) 0%, rgba(118, 75, 162, 0.15) 100%);
        }
    }
    
    /* Text areas dark mode support */
    @media (prefers-color-scheme: dark) {
        textarea {
            background-color: #1E1E1E !important;
            color: #FFFFFF !important;
            border-color: #404040 !important;
        }
        .stTextArea textarea {
            background-color: #1E1E1E !important;
            color: #FFFFFF !important;
        }
        .stTextInput input {
            background-color: #1E1E1E !important;
            color: #FFFFFF !important;
        }
    }

    /* Style text inputs to match the dark rounded pill from the screenshot */
    .stTextInput > div > div > input {
        background-color: #2f2f2f !important;
        color: #e6e6e6 !important;
        height: 40px !important;
        font-size: 1rem !important;
        line-height: 1.4 !important;
        width: 100% !important;
        max-width: 100% !important;
        box-shadow: 0 8px 24px rgba(0,0,0,0.6) inset, 0 2px 8px rgba(0,0,0,0.6) !important;
    }
    
    .stTextInput > div > div > input:hover {
        background-color: #353535 !important;
    }
    
    .stTextInput > div > div > input:focus {
        outline: none !important;
        box-shadow: 0 8px 24px rgba(0,0,0,0.6) inset, 0 0 0 2px #555 !important;
    }
    
    /* Style the submit button - perfectly aligned with input field */
    div[data-testid="column"]:last-child button {
        background-color: white !important;
        color: #333 !important;
        border: 2px solid #e0e0e0 !important;
        border-radius: 12px !important;
        width: 65px !important;
        height: 65px !important;
        padding: 0 !important;
        font-size: 24px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1) !important;
        transition: all 0.3s ease !important;
        margin: 0 !important;
        position: relative !important;
        top: 0 !important;
    }
    
    div[data-testid="column"]:last-child button:hover {
        background-color: #f8f8f8 !important;
        border-color: #ccc !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15) !important;
    }
    
    div[data-testid="column"]:last-child button:active {
        transform: translateY(0px) !important;
        box-shadow: 0 1px 4px rgba(0,0,0,0.1) !important;
    }
    
    /* Perfectly align button column container with input */
    div[data-testid="column"]:last-child div[data-testid="stButton"] {
        display: flex !important;
        align-items: flex-start !important;
        margin-top: 0 !important;
        padding-top: 0 !important;
        height: 65px !important;
        justify-content: center !important;
    }
    
    /* Ensure button container has no extra spacing */
    div[data-testid="column"]:last-child {
        display: flex !important;
        align-items: flex-start !important;
        padding-top: 0 !important;
        margin-top: 0 !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Sidebar removed per user request (logos and sidebar controls hidden)

# Extract page name without emoji
# Navigation removed: default to Home
page = "Home"
page_name = page.split(" ", 1)[1] if " " in page else page

if page_name == "Home":
    # V7-style hero section with description - minimal top spacing
    st.markdown("""
    <div style="text-align: center; padding: 0.5rem 2rem 2rem 2rem; max-width: 1200px; margin: 0 auto;">
        <h1 style="font-size: 4rem; font-weight: 800; line-height: 1.2; margin-bottom: 1rem; margin-top: 0;">
            <span style="color: #2D3748;">Analyze your Contracts</span><br>
            <span style="color: #A0AEC0;">with</span>
            <span style="color: #1E8C7E;"> Naja7</span>
        </h1>
        <p style="font-size: 1.3rem; color: #718096; max-width: 800px; margin: 0 auto 2rem auto; line-height: 1.6;">
            Upload your PDF and Word documents to extract, analyze, and understand contract content with AI-powered insights. 
            Ask questions about your contracts, extract key information, and streamline your document review process with intelligent automation.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Centered container for upload
    col1, col2, col3 = st.columns([1, 3, 1])
    with col2:
        # File drop area (drag-and-drop) — placed above the Q&A textbox per user request
        home_uploaded_files = st.file_uploader(
            "Drop your PDF or Word files here or click to browse",
            type=["pdf", "docx", "doc"],
            help="Supported formats: PDF (.pdf) and Word (.docx)",
            accept_multiple_files=True,
            key="home_uploader"
        )

        # Q&A header (input removed per user request)
        st.markdown("### Spcify the country/region (optional)")

        # Show uploaded files as chips (like V7)
        if home_uploaded_files:
            st.markdown("<div style='margin-top: 1rem;'>", unsafe_allow_html=True)
            file_chips = ""
            for file in home_uploaded_files:
                file_ext = file.name.split('.')[-1].lower()
                icon = "📄" if file_ext == "pdf" else "📝"
                file_chips += f"""
                <span style="display: inline-block; background: #F7FAFC; border: 1px solid #E2E8F0; 
                             padding: 0.5rem 1rem; border-radius: 20px; margin: 0.25rem; 
                             font-size: 0.9rem; color: #2D3748;">
                    {icon} {file.name}
                </span>
                """
            st.markdown(file_chips, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # Add a compact rounded input with button on the right
        input_col, button_col = st.columns([6, 1])
        
        with input_col:
            qa_followup = st.text_input(
                "Q&A Input",
                placeholder="Ask anything",
                key="qa_followup_input",
                label_visibility="collapsed",
                help="Ask anything about your documents"
            )
        
        with button_col:
            # Perfect alignment with input field
            if st.button("⏹", key="submit_qa", help="Generate response"):
                if qa_followup:
                    # Generate text action with custom styled message
                    st.markdown("""
                    <div style="
                        width: 150px; 
                        background-color: #d4edda; 
                        color: #155724; 
                        border: 1px solid #c3e6cb; 
                        border-radius: 8px; 
                        padding: 8px 12px; 
                        font-size: 14px; 
                        margin-top: 8px;
                        text-align: center;
                    ">
                        Generating response...
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="
                        width: 150px; 
                        background-color: #fff3cd; 
                        color: #856404; 
                        border: 1px solid #ffeaa7; 
                        border-radius: 8px; 
                        padding: 8px 12px; 
                        font-size: 14px; 
                        margin-top: 8px;
                        text-align: center;
                    ">
                        Please enter a question first
                    </div>
                    """, unsafe_allow_html=True)
    
    # Privacy Information Cards Section
    st.markdown("<br><br>", unsafe_allow_html=True)  # Add spacing between sections
    
    # Add title for the cards section
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2rem;">
        <h2 style="
            font-size: 2.5rem; 
            font-weight: 700; 
            color: #2D3748; 
            margin-bottom: 0.5rem;
            letter-spacing: -1px;
        ">Privacy & Security</h2>
        <p style="
            font-size: 1.1rem; 
            color: #718096; 
            margin: 0 auto; 
            max-width: 600px;
            line-height: 1.5;
        ">Learn how we protect your documents and ensure your data privacy throughout the analysis process.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Create 3 columns for the privacy cards
    card_col1, card_col2, card_col3 = st.columns(3)
    
    with card_col1:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
            padding: 2rem; 
            border-radius: 15px; 
            color: white; 
            text-align: center;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <h3 style="color: white; margin-bottom: 1rem;">How Privacy Works</h3>
            <p style="font-size: 0.9rem; line-height: 1.5; opacity: 0.9;">
                Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with card_col2:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); 
            padding: 2rem; 
            border-radius: 15px; 
            color: white; 
            text-align: center;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <h3 style="color: white; margin-bottom: 1rem;">How Privacy Works</h3>
            <p style="font-size: 0.9rem; line-height: 1.5; opacity: 0.9;">
                Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with card_col3:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); 
            padding: 2rem; 
            border-radius: 15px; 
            color: white; 
            text-align: center;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <h3 style="color: white; margin-bottom: 1rem;">How Privacy Works</h3>
            <p style="font-size: 0.9rem; line-height: 1.5; opacity: 0.9;">
                Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)  # Add spacing after cards
    
    if home_uploaded_files:
        st.success(f"✅ {len(home_uploaded_files)} file(s) uploaded successfully!")
        
        # Process files
        all_texts = []
        for uploaded_file in home_uploaded_files:
            try:
                file_name = uploaded_file.name
                file_type = uploaded_file.type
                
                with st.spinner(f'🔄 Processing {file_name}...'):
                    full_text = ""
                    num_pages = 0
                    
                    if "pdf" in file_type.lower() or file_name.lower().endswith('.pdf'):
                        pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
                        num_pages = len(pdf_reader.pages)
                        for page in pdf_reader.pages:
                            full_text += page.extract_text() + "\n"
                    elif file_name.lower().endswith('.docx') or file_name.lower().endswith('.doc'):
                        doc = Document(BytesIO(uploaded_file.read()))
                        num_pages = len(doc.paragraphs) // 20 + 1
                        for paragraph in doc.paragraphs:
                            full_text += paragraph.text + "\n"
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    full_text += cell.text + " "
                            full_text += "\n"
                    
                    all_texts.append({
                        'filename': file_name,
                        'text': full_text,
                        'pages': num_pages
                    })
                    
                st.success(f"✅ {file_name} processed - {num_pages} pages")
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")
        
        # Output box
        st.markdown("### 📤 Extracted Text")
        st.markdown('<div class="output-box">', unsafe_allow_html=True)
        combined_text = "\n\n".join([f"=== {item['filename']} ===\n{item['text']}" for item in all_texts])
        st.text_area("All Extracted Text", combined_text, height=300, label_visibility="collapsed")
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Download button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.download_button(
                label="📥 Download Results",
                data=combined_text,
                file_name="extracted_text.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        # Q&A Section
        st.markdown("### 💬 Questions & Answers")
        st.markdown('<div class="qa-box">', unsafe_allow_html=True)
        
        if 'qa_pairs' not in st.session_state:
            st.session_state.qa_pairs = []
        
        question = st.text_input("Your Question", placeholder="Ask about your Contracts...")
        answer = st.text_area("Answer", placeholder="Type your answer...", height=100)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save Q&A", use_container_width=True):
                if question and answer:
                    st.session_state.qa_pairs.append({'question': question, 'answer': answer})
                    st.success("✅ Q&A saved!")
        with col2:
            if st.button("🗑️ Clear History", use_container_width=True):
                st.session_state.qa_pairs = []
                st.success("✅ History cleared!")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Show Q&A history
        if st.session_state.qa_pairs:
            st.markdown("### 📋 Q&A History")
            st.markdown('<div class="output-box">', unsafe_allow_html=True)
            qa_output = "\n\n".join([f"Q: {qa['question']}\nA: {qa['answer']}" for qa in st.session_state.qa_pairs])
            st.text_area("Q&A Output", qa_output, height=200, label_visibility="collapsed")
            st.markdown('</div>', unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="📥 Download Q&A",
                    data=qa_output,
                    file_name="qa_history.txt",
                    mime="text/plain",
                    use_container_width=True
                )

elif page_name == "Document Analyzer":
    st.markdown('<h1 class="main-header">📄 Document Analyzer</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Upload your PDF or Word document to begin</p>', unsafe_allow_html=True)
    
    # Multiple file upload with drag and drop
    st.markdown("### 📁 Upload Contracts")
    uploaded_files = st.file_uploader(
        "Drop your files here or click to browse (supports multiple files)",
        type=["pdf", "docx", "doc"],
        help="Supported formats: PDF (.pdf) and Word (.docx)",
        accept_multiple_files=True
    )
    
    # Show code implementation dropdown
    with st.expander("💻 View Implementation Code"):
        st.code("""
# Import required libraries
import streamlit as st
import PyPDF2
from docx import Document
from io import BytesIO

# Multiple file upload
uploaded_files = st.file_uploader(
    "Upload Contracts",
    type=["pdf", "docx", "doc"],
    accept_multiple_files=True
)

# Process each file
for uploaded_file in uploaded_files:
    if uploaded_file.name.endswith('.pdf'):
        # Process PDF
        pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\\n"
    
    elif uploaded_file.name.endswith('.docx'):
        # Process Word document
        doc = Document(BytesIO(uploaded_file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\\n"
        """, language="python")
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} file(s) uploaded successfully!")
        
        # Process all files
        all_texts = []
        for uploaded_file in uploaded_files:
            try:
                file_type = uploaded_file.type
                file_name = uploaded_file.name
                file_size = uploaded_file.size / 1024  # Convert to KB
                
                # Display file info in an attractive card
                st.markdown(f"### 📋 Processing: {file_name}")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                                padding: 1.5rem; border-radius: 15px; color: white; text-align: center;">
                        <h4 style="margin: 0; color: white;">📄 File Name</h4>
                        <p style="margin-top: 0.5rem; font-size: 0.9rem;">{file_name}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    file_extension = file_name.split('.')[-1].upper()
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); 
                                padding: 1.5rem; border-radius: 15px; color: white; text-align: center;">
                        <h4 style="margin: 0; color: white;">📋 File Type</h4>
                        <p style="margin-top: 0.5rem; font-size: 0.9rem;">{file_extension}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col3:
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); 
                                padding: 1.5rem; border-radius: 15px; color: white; text-align: center;">
                        <h4 style="margin: 0; color: white;">💾 File Size</h4>
                        <p style="margin-top: 0.5rem; font-size: 0.9rem;">{file_size:.2f} KB</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                full_text = ""
                num_pages = 0
                
                # Processing indicator
                with st.spinner('🔄 Processing your document...'):
                    # Read based on file type
                    if "pdf" in file_type.lower() or file_name.lower().endswith('.pdf'):
                        pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
                        num_pages = len(pdf_reader.pages)
                        for page in pdf_reader.pages:
                            full_text += page.extract_text() + "\n"
                            
                    elif file_name.lower().endswith('.docx') or file_name.lower().endswith('.doc'):
                        doc = Document(BytesIO(uploaded_file.read()))
                        num_pages = len(doc.paragraphs) // 20 + 1
                        for paragraph in doc.paragraphs:
                            full_text += paragraph.text + "\n"
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    full_text += cell.text + " "
                            full_text += "\n"
                
                all_texts.append({
                    'filename': file_name,
                    'text': full_text,
                    'pages': num_pages
                })
                
                st.success(f"✅ {file_name} processed successfully! {num_pages} pages analyzed.")
                st.markdown("---")
                
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")
                continue
        
        # Combined results output box
        st.markdown("### 📤 Analysis Output")
        st.markdown('<div class="output-box">', unsafe_allow_html=True)
        
        combined_text = "\n\n".join([f"=== {item['filename']} ===\n{item['text']}" for item in all_texts])
        output_text = st.text_area(
            "All Extracted Text",
            combined_text,
            height=300,
            help="Combined text from all uploaded Contracts"
        )
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Download combined results
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.download_button(
                label="📥 Download All Results",
                data=combined_text,
                file_name="combined_analysis.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        # Q&A Section
        st.markdown("### 💬 Questions & Answers")
        st.markdown('<div class="qa-box">', unsafe_allow_html=True)
        
        st.markdown("#### Ask questions about your Contracts:")
        
        # Initialize Q&A in session state
        if 'qa_pairs' not in st.session_state:
            st.session_state.qa_pairs = []
        
        # Question input removed per user request
        question = ""
        
        col1, col2 = st.columns([3, 1])
        with col1:
            answer = st.text_area(
                "Answer",
                placeholder="Type your answer here or let AI generate it...",
                height=100,
                key="answer_input"
            )
        
        with col2:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("💾 Save Q&A", use_container_width=True):
                if question and answer:
                    st.session_state.qa_pairs.append({
                        'question': question,
                        'answer': answer
                    })
                    st.success("✅ Q&A saved!")
                else:
                    st.warning("⚠️ Please enter both question and answer")
            
            if st.button("🤖 Auto-Answer", use_container_width=True):
                if question and combined_text:
                    # Simple keyword-based answer (you can replace with AI later)
                    question_lower = question.lower()
                    relevant_text = ""
                    
                    # Find sentences containing question keywords
                    keywords = [word for word in question_lower.split() if len(word) > 3] # type: ignore
                    sentences = re.split(r'[.!?]+', combined_text)
                    
                    for sentence in sentences[:50]:  # Check first 50 sentences
                        if any(keyword in sentence.lower() for keyword in keywords):
                            relevant_text += sentence.strip() + ". "
                    
                    if relevant_text:
                        st.session_state.temp_answer = relevant_text[:500] + "..."
                        st.info("📝 Auto-generated answer based on document content")
            st.markdown("### 📋 Saved Q&A History")
            st.markdown('<div class="output-box">', unsafe_allow_html=True)
            
            qa_output = ""
            for i, qa in enumerate(st.session_state.qa_pairs, 1):
                qa_output += f"Q{i}: {qa['question']}\n"
                qa_output += f"A{i}: {qa['answer']}\n\n"
            
            st.text_area(
                "Q&A Output",
                qa_output,
                height=200,
                help="All your saved questions and answers"
            )
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Download Q&A
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="📥 Download Q&A",
                    data=qa_output,
                    file_name="qa_history.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            if st.button("🗑️ Clear Q&A History"):
                st.session_state.qa_pairs = []
                st.rerun()
        
        # Statistics tabs
        st.markdown("### 📊 Detailed Analysis")
        tab1, tab2, tab3, tab4 = st.tabs([
            "📖 Text Content",
            "📊 Statistics",
            "🔍 Search",
            "📈 Word Analysis"
        ])
        
        with tab1:
            st.markdown("### 📖 Individual Contracts")
            for item in all_texts:
                with st.expander(f"📄 {item['filename']} ({item['pages']} pages)"):
                    st.text_area(
                        f"Content of {item['filename']}",
                        item['text'],
                        height=300,
                        label_visibility="collapsed"
                    )
        
        with tab2:
            st.markdown("### 📊 Combined Statistics")
            
            # Calculate combined statistics
            total_word_count = sum(len(item['text'].split()) for item in all_texts)
            total_char_count = sum(len(item['text']) for item in all_texts)
            total_pages = sum(item['pages'] for item in all_texts)
            
            all_sentences = []
            for item in all_texts:
                sentences = re.split(r'[.!?]+', item['text'])
                all_sentences.extend([s.strip() for s in sentences if s.strip()])
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                            padding: 2rem; border-radius: 15px; color: white; text-align: center;
                            box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <h1 style="margin: 0; color: white; font-size: 2.5rem;">{total_pages}</h1>
                    <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">📄 Total Pages</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); 
                            padding: 2rem; border-radius: 15px; color: white; text-align: center;
                            box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <h1 style="margin: 0; color: white; font-size: 2.5rem;">{total_word_count:,}</h1>
                    <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">📝 Total Words</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); 
                            padding: 2rem; border-radius: 15px; color: white; text-align: center;
                            box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <h1 style="margin: 0; color: white; font-size: 2.5rem;">{total_char_count:,}</h1>
                    <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">🔤 Total Characters</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col4:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #43e97b 0%, #38f9d7 100%); 
                            padding: 2rem; border-radius: 15px; color: white; text-align: center;
                            box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <h1 style="margin: 0; color: white; font-size: 2.5rem;">{len(all_sentences):,}</h1>
                    <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">📜 Total Sentences</p>
                </div>
                """, unsafe_allow_html=True)
        
        with tab3:
            st.markdown("### 🔍 Search Across All Contracts")
            search_term = st.text_input("Enter search term", placeholder="e.g., contract, policy, agreement")
            
            if search_term:
                st.markdown("#### Search Results:")
                total_occurrences = 0
                
                for item in all_texts:
                    count = item['text'].lower().count(search_term.lower())
                    total_occurrences += count
                    
                    if count > 0:
                        st.markdown(f"**{item['filename']}**: Found {count} occurrence(s)")
                        
                        # Show context
                        lines = item['text'].split('\n')
                        shown = 0
                        for i, line in enumerate(lines):
                            if search_term.lower() in line.lower() and shown < 3:
                                st.text(f"Line {i+1}: {line[:200]}...")
                                shown += 1
                
                st.info(f"📊 Total occurrences: **{total_occurrences}**")
        
        with tab4:
            st.markdown("### 📈 Word Frequency Analysis")
            
            # Combine all text for word frequency
            all_words = []
            for item in all_texts:
                words = item['text'].lower().split()
                words = [re.sub(r'[^\w\s]', '', word) for word in words if len(word) > 3]
                all_words.extend(words)
            
            if all_words:
                word_freq = {}
                for word in all_words:
                    word_freq[word] = word_freq.get(word, 0) + 1
                
                # Sort by frequency
                sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
                
                # Create dataframe
                df = pd.DataFrame(sorted_words, columns=['Word', 'Frequency'])
                
                st.markdown("#### Top 20 Most Common Words")
                st.dataframe(df, use_container_width=True)
                
                st.bar_chart(df.set_index('Word'))
    
    else:
        # Beautiful upload prompt
        st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    padding: 4rem; border-radius: 20px; text-align: center; color: white; margin: 2rem 0;">
            <h1 style="color: white; font-size: 3rem; margin: 0;">📁</h1>
            <h2 style="color: white; margin: 1rem 0;">Ready to Analyze?</h2>
            <p style="font-size: 1.2rem; opacity: 0.9; margin: 0;">Upload one or more PDF or Word Contracts to get started</p>
        </div>
        """, unsafe_allow_html=True)

# Footer
# Sidebar removed; no footer in sidebar

